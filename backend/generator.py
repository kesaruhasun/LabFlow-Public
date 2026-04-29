import nbformat as nbf
from nbclient import NotebookClient
from nbconvert import HTMLExporter
from playwright.async_api import async_playwright
import os
import subprocess
import time
import asyncio
import textwrap
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

async def generate_submission_document(work_dir: str, questions: list, it_number: str, name: str, lab_no: str) -> str:
    notebook_path = os.path.join(work_dir, "submission_notebook.ipynb")
    html_path = os.path.join(work_dir, "temp_notebook.html")
    docx_path = os.path.join(work_dir, "final_submission.docx")
    
    # 1. Generate Notebook with Input Mocking
    print("Generating Jupyter Notebook Workspace...")
    nb = nbf.v4.new_notebook()
    
    # Prepend an input mocker so headless execution doesn't crash on input() calls
    mocker_code = textwrap.dedent("""
    import builtins
    import random
    import sys
    
    _input_call_count = 0
    _MAX_INPUT_CALLS = 150
    
    def mock_input(prompt=""):
        global _input_call_count
        _input_call_count += 1
        
        if _input_call_count > _MAX_INPUT_CALLS:
            print(f"{prompt} [MAX INPUT CALLS REACHED. BREAKING INFINITE LOOP]")
            raise Exception("Infinite loop detected: input() overcalled across the entire notebook.")
            
        prompt_lower = prompt.lower()
        
        # Intelligent Type Inference
        if 'password' in prompt_lower:
            val = 'python123'
        elif 'name' in prompt_lower:
            val = 'Alice'
        elif 'str' in prompt_lower or 'word' in prompt_lower or 'text' in prompt_lower or 'string' in prompt_lower:
            val = 'example_string'
        elif 'float' in prompt_lower or 'decimal' in prompt_lower or 'weight' in prompt_lower or 'temperature' in prompt_lower or 'average' in prompt_lower:
            val = '42.5'
        elif 'int' in prompt_lower or 'mark' in prompt_lower or 'score' in prompt_lower or 'age' in prompt_lower or 'count' in prompt_lower or 'number' in prompt_lower:
            val = str(random.randint(10, 99))
        else:
            # Fallback to integer for ambiguous prompts since many intro labs cast inputs to int/float
            val = str(random.randint(5, 50))
            
        print(f"{prompt}{val}")
        return val
        
    # Override in builtins
    builtins.input = mock_input
    
    # Also override in global namespace to shadow kernel builtins
    globals()['input'] = mock_input
    
    # Monkey-patch IPython kernel's raw_input just in case
    try:
        from ipykernel.kernelbase import Kernel
        Kernel.raw_input = lambda self, prompt="": mock_input(prompt)
    except ImportError:
        pass
    """)
    nb.cells.append(nbf.v4.new_code_cell(mocker_code.strip()))
    
    for q in questions:
        # Assuming AI generated a 'code' string for each question
        nb.cells.append(nbf.v4.new_code_cell(q["code"]))
    
    # 2. Execute Notebook securely
    print("Executing Notebook...")
    client = NotebookClient(nb, timeout=600, kernel_name='python3')
    await client.async_execute()

    with open(notebook_path, "w") as f:
        nbf.write(nb, f)

    # 3. Convert to Authentic HTML
    print("Converting to HTML with typical JupyterLab styling...")
    html_exporter = HTMLExporter(template_name='lab')
    (body, resources) = html_exporter.from_notebook_node(nb)
    
    style = """
    <style>
    body { padding: 40px; background-color: #ffffff; font-family: 'Times New Roman', Times, serif; }
    .jp-Cell { margin-bottom: 20px !important; border: 1px solid #e0e0e0 !important; border-radius: 4px; padding: 10px; }
    .jp-OutputArea-output { pre { white-space: pre-wrap; font-family: monospace; } }
    .student-header { text-align: center; margin-bottom: 50px; }
    .student-header h2 { margin: 5px 0; color: #111; font-size: 24px; font-weight: normal; }
    .student-header h3 { margin: 5px 0; color: #111; font-size: 18px; font-weight: normal; }
    .student-header h4 { margin: 15px 0 5px 0; color: #111; font-size: 16px; font-weight: normal; }
    .page-break { page-break-after: always; }
    </style>
    """
    
    header_html = f"""
    <div class="student-header">
        <h2>University / Institution Name</h2>
        <h3>Lab Submission</h3>
        <h3>{lab_no}</h3>
        <h3>{it_number}</h3>
        <h3>{name}</h3>
        <h4>Course Module Name</h4>
        <h4>Degree Program Name</h4>
    </div>
    <div class="page-break"></div>
    """
    
    if "<body>" in body:
        body = body.replace("<body>", f"<body>{style}{header_html}")
    else:
        body += style + header_html

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(body)

    # 4. Screenshot the cells via Playwright
    print("Taking screenshots of each cell...")
    screenshot_paths = []
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True, args=['--no-sandbox', '--disable-setuid-sandbox', '--disable-dev-shm-usage'])
        page = await browser.new_page(viewport={'width': 1000, 'height': 800})
        await page.goto(f"file://{os.path.abspath(html_path)}")
        await page.wait_for_load_state("networkidle")
        await asyncio.sleep(2) 
        
        cells = page.locator('.jp-Cell')
        count = await cells.count()
        # Specifically skip the first cell (index 0) because it contains the hidden input() mocker
        for i in range(1, count):
            img_path = os.path.join(work_dir, f"q_{i}.png")
            await cells.nth(i).screenshot(path=img_path)
            screenshot_paths.append(img_path)
            
        await browser.close()

    # 5. Assemble to Document
    print("Generating Document from Template...")
    # Read the template from the local backend folder
    template_path = os.path.join(os.path.dirname(__file__), 'Generic_Lab_Template.docx')
    doc = Document(template_path)
    
    # Replace template placeholders completely with specific styling
    for p in doc.paragraphs:
        if '<Lab sheet No>' in p.text:
            new_text = p.text.replace('<Lab sheet No>', lab_no).strip()
            p.text = "" # Clear original placeholder text
            run = p.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(26)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if '<ITXXXXXXXX>' in p.text:
            new_text = p.text.replace('<ITXXXXXXXX>', it_number).strip()
            p.text = ""
            run = p.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(26)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if '<Name>' in p.text:
            new_text = p.text.replace('<Name>', name).strip()
            p.text = ""
            run = p.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(26)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    for idx, (q, img_path) in enumerate(zip(questions, screenshot_paths)):
        p = doc.add_paragraph()
        run = p.add_run(q.get('heading', f"Question {idx+1}"))
        run.bold = True
        # Deliberately not setting font.size so it inherits the default template styles (Times New Roman)
        
        doc.add_picture(img_path, width=Inches(6.0))
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(24)

    doc.save(docx_path)
    
    # 6. Convert Docx to true PDF using LibreOffice Headless Subprocess
    answer_sheet_pdf_path = os.path.join(work_dir, "answer_sheet.pdf")
    try:
        # LibreOffice outputs to the outdir with the original filename but .pdf extension
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', work_dir],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        # Rename the output to our expected answer_sheet.pdf
        generated_pdf = docx_path.replace('.docx', '.pdf')
        if os.path.exists(generated_pdf):
            os.rename(generated_pdf, answer_sheet_pdf_path)
    except Exception as e:
        print(f"LibreOffice PDF conversion failed: {e}")
        # Fallback empty string if conversion entirely fails (should not happen if container has libreoffice)
        pass 
        
    return docx_path, answer_sheet_pdf_path
